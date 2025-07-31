import os
from docx import *
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.shared import Cm
from docxcompose.composer import Composer
import sys
import json
from docx2pdf import convert


sys.stdout.reconfigure(encoding='utf-8')  # ✅ 保證輸出為 UTF-8（避免 Windows 預設 gbk/cp950）



def fillin(userName , chooseFile , inputJsons , fileType):
    wordFilePaths = []
    for inputJson in inputJsons:
        inputJson = inputJson[0]
        wordPath = './convert_content/5.報名表正面.docx'
        wordDoc = Document(wordPath)
        table = wordDoc.tables[0]
        filledInfo = set()
        for idxr, row in enumerate(table.rows):
            for idxc , cell in enumerate(row.cells):
                if cell.text == "" and table.cell(idxr,idxc-1) not in filledInfo:
                    fillitemIdx = table.cell(idxr,idxc-1).text.replace('\n','')
                    try:
                        cell.text = inputJson[fillitemIdx]
                        filledInfo.add(inputJson[fillitemIdx])
                    except:
                        pass
                elif table.cell(idxr,idxc-1).text.replace("\n" , '') == "報檢職類":
                    type_dict = {
                        '視覺':'視覺傳達設計',
                        '會計人工':'會計事務 -人工記帳',
                        '會計資訊':'會計事務 -資訊',
                        '會資':'會計事務 -資訊',
                        '門市':'門市服務'
                    }
                    fillInWd = cell.text
                    fillLoc = fillInWd.index(type_dict[inputJson['報簡職類']])
                    aboveWd = fillInWd[:fillLoc-1]
                    behindWd = fillInWd[fillLoc:]
                    fiWd = aboveWd + "■" + behindWd
                    fillLoc = fiWd.index('網頁設計')
                    aboveWd  = fiWd[:fillLoc-2]
                    behindWd = fiWd[fillLoc-1:]
                    fiWd = aboveWd + behindWd
                    fillLoc = fiWd.index('視覺傳達設計')
                    aboveWd  = fiWd[:fillLoc-3]
                    behindWd = fiWd[fillLoc-1:]
                    fiWd = aboveWd + behindWd
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.clear()
                    paragraph  = cell.paragraphs[0]
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = Pt(14)
                    run = paragraph.add_run(fiWd)
                    run.font.size = Pt(11.5)
                    run.font.name = 'Times New Roman' 
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    spacing = OxmlElement('w:spacing')
                    spacing.set(qn('w:val'), '0')  
                elif table.cell(idxr,idxc-1).text.replace("\n" , '') == "英文姓名":
                    paragraphs = cell.paragraphs
                    paragraphs[0].text  = inputJson['英文姓名']
                elif table.cell(idxr,idxc-1).text.replace("\n" , '') == "檢定區別" :
                    testTypeDic = {
                        "全測":"學術科全測 ",
                        "免學":"A免試學科",
                        "免術":"B免試術科"
                    }
                    if cell.text  in filledInfo:
                        continue
                    paragraphs = cell.paragraphs
                    for idx , paragraph in enumerate(paragraphs):
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = Pt(8)
                        if idx ==0:
                            finalText = "□" + paragraph.text[1:]
                            paragraph.clear()
                            run = paragraph.add_run(finalText)
                            run.font.size = Pt(11.2)
                            run.font.name = 'Times New Roman' 
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                            rPr = run._element.get_or_add_rPr()
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:val'), '0.5')  
                            rPr.append(spacing)
                        else:
                            finalText = "□" + paragraph.text[1:]
                            paragraph.clear()
                            aboveWord = finalText[:6]
                            behindWord = finalText[6:]
                            run =paragraph.add_run(aboveWord)
                            run.font.size = Pt(11.2)
                            run.font.name = 'Times New Roman' 
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                            run = paragraph.add_run(behindWord)
                            run.font.size = Pt(8)
                            run.font.name = 'Times New Roman' 
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                            rPr = run._element.get_or_add_rPr()
                            spacing = OxmlElement('w:spacing')
                            spacing.set(qn('w:val'), '0.5')  
                            rPr.append(spacing)
                    for paragraph in paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = Pt(8)
                        if testTypeDic[inputJson['檢定區別']] in paragraph.text:
                            if cell.text in filledInfo:
                                continue
                            if testTypeDic[inputJson['檢定區別']] == "學術科全測":
                                finalText = "■" + paragraph.text[1:]
                                paragraph.clear()
                                run = paragraph.add_run(finalText)
                                run.font.size = Pt(11.2)
                                run.font.name = 'Times New Roman' 
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                rPr = run._element.get_or_add_rPr()
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:val'), '0.5')  
                                rPr.append(spacing)
                            else:        
                                finalText = "■" + paragraph.text[1:]
                                paragraph.clear()
                                aboveWord = finalText[:6]
                                behindWord = finalText[6:]
                                run =paragraph.add_run(aboveWord)
                                run.font.size = Pt(11.2)
                                run.font.name = 'Times New Roman' 
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                run = paragraph.add_run(behindWord)
                                run.font.size = Pt(7.8)
                                run.font.name = 'Times New Roman' 
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                                rPr = run._element.get_or_add_rPr()
                                spacing = OxmlElement('w:spacing')
                                spacing.set(qn('w:val'), '0.5')  
                                rPr.append(spacing)
                    filledInfo.add(cell.text)
                elif table.cell(idxr,idxc-1).text.replace("\n" , '') == "聯絡電話" :
                    if cell.text.replace("\n" , "")  in filledInfo:
                        # print("im jump out")
                        continue
                    paragraphs =  cell.paragraphs
                    for paragraph in paragraphs:
                        if "住宅" in paragraph.text:
                            paragraph.text += inputJson['聯絡電話(住宅)']
                            runs = paragraph.runs
                            for run in runs:
                                run.font.size = Pt(12)
                                run.font.name = "Times New Roman"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

                        elif "手機" in paragraph.text:
                            paragraph.text += inputJson['聯絡電話(手機)']
                            runs = paragraph.runs
                            for run in runs:
                                run.font.size = Pt(12)
                                run.font.name = "Times New Roman"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    filledInfo.add(cell.text.replace("\n" , ""))
                    # print(cell.text.replace("\n" , ""))
                elif  "身分別" in  table.cell(idxr,idxc-1).text.replace("\n" , '') :
                    paragraphs = cell.paragraphs
                    fillin = inputJson['身分別']
                    # 先做資料格式設定
                    for paragraph in paragraphs :
                        paragraph.paragraph_format.line_spacing = Pt(8.5)
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                    for paragraph in paragraphs:
                        if inputJson['身分別'] in paragraph.text:
                            if cell.text in filledInfo:
                                continue
                            findWdLoc = paragraph.text.index(inputJson['身分別'])
                            aboveWd = paragraph.text[:findWdLoc-3]
                            behindWd = "■" + paragraph.text[findWdLoc-2:]
                            paragraph.text = aboveWd + behindWd
                            for run in paragraph.runs:
                                run.font.size = Pt(8)
                                run.font.name = "Times New Roman"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    for paragraph in paragraphs:
                        newString= ""
                        for wd in paragraph.text:
                            if wd =="□":
                                newString += "□"
                            else:
                                newString += wd
                        paragraph.text = newString
                        for run in paragraph.runs:
                            run.font.size = Pt(8)
                            run.font.name = "Times New Roman"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    filledInfo.add(cell.text)
                elif cell.text  == "上課別":
                    targetCell  = table.cell(idxr+1 , idxc)
                    paragraphs = targetCell.paragraphs
                    paragraphs[0].text = inputJson['上課別']
                elif cell.text  == "年級":
                    targetCell  = table.cell(idxr+1 , idxc)
                    paragraphs = targetCell.paragraphs
                    paragraphs[0].text = inputJson["年級"]
                elif cell.text  == "班別":
                    targetCell  = table.cell(idxr+1 , idxc)
                    paragraphs = targetCell.paragraphs
                    paragraphs[0].text = inputJson['班級']
                elif cell.text  == "座號":
                    targetCell  = table.cell(idxr+1 , idxc)
                    paragraphs = targetCell.paragraphs
                    paragraphs[0].text = inputJson['座號']
                elif  "學制" in  table.cell(idxr,idxc-1).text.replace("\n" , '') :
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs :
                        paragraph.paragraph_format.line_spacing = Pt(12)
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                    for paragraph in paragraphs:
                        if inputJson['學制'] in paragraph.text:
                            if cell.text in filledInfo:
                                continue
                            findWdLoc = paragraph.text.index(inputJson['學制'])
                            aboveWd = paragraph.text[:findWdLoc-1]
                            behindWd = "■" + paragraph.text[findWdLoc:]
                            paragraph.text = aboveWd + behindWd
                            for run in paragraph.runs:
                                run.font.size = Pt(12)
                                run.font.name = "Times New Roman"
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    for paragraph in paragraphs:
                        newString= ""
                        for wd in paragraph.text:
                            if wd =="□":
                                newString += "□"
                            else:
                                newString += wd
                        paragraph.text = newString
                        for run in paragraph.runs:
                            run.font.size = Pt(12)
                            run.font.name = "Times New Roman"
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
                    filledInfo.add(cell.text)
                elif  "實貼身分證【正面】" in  table.cell(idxr,idxc-1).text.replace("\n" , '') and cell.text =="    ":
                    paragraphs = cell.paragraphs
                    for paragraph in paragraphs:
                        paragraph.clear()
                    
                    paragraph = cell.add_paragraph()
                    paragraph.alignment  = WD_ALIGN_PARAGRAPH.CENTER
                    run = paragraph.add_run()
                    run.add_picture(f'./user_data/{userName}/{chooseFile}/image/{inputJson["身分證號碼"]}.jpg')
                    run = paragraph.add_run()
                    run.add_picture(f'./user_data/{userName}/{chooseFile}/image/{inputJson["身分證號碼"]}.jpg')
        wordDoc.save(f'./user_data/{userName}/{chooseFile}/{fileType}/{inputJson["身分證號碼"]}.docx')
        wordFilePaths.append(f'./user_data/{userName}/{chooseFile}/{fileType}/{inputJson["身分證號碼"]}.docx')
    return wordFilePaths

def combineWord(wordFilePaths , fileType):
    # print(wordFilePaths)
    master = Document(wordFilePaths[0])
    composer = Composer(master)
    for file in wordFilePaths[1:]:
        doc = Document(file)
        composer.append(doc)
    composer.save(f'./user_data/{userName}/{chooseFile}/{fileType}/combine.docx')
    return f'./user_data/{userName}/{chooseFile}/{fileType}/combine.docx'

def docxConvert():
    convert(f'./user_data/{userName}/{chooseFile}/combine.docx',f'./user_data/{userName}/{chooseFile}/combine.pdf')

if __name__ == "__main__":
    print("im running python")
    userName = sys.argv[1]
    chooseFile = sys.argv[2]
    combineWordPath = []
    if os.path.exists(f'./user_data/{userName}/{chooseFile}/fullTest/fullTest.json'):
        with open(f'./user_data/{userName}/{chooseFile}/fullTest/fullTest.json' , encoding="utf-8") as f:
            inputJsons = json.load(f) 
        wordFilePath = fillin(userName , chooseFile , inputJsons  , "fullTest")
        path = combineWord(wordFilePath , "fullTest")
        combineWordPath.append(path)
        # docxConvert()

    if os.path.exists(f'./user_data/{userName}/{chooseFile}/studyTest/studyTest.json'):
        with open(f'./user_data/{userName}/{chooseFile}/studyTest/studyTest.json' , encoding="utf-8") as f:
            inputJsons = json.load(f)
        wordFilePath = fillin(userName , chooseFile , inputJsons , "studyTest")
        path = combineWord(wordFilePath , "studyTest")  
        combineWordPath.append(path) 
        # docxConvert()

    if os.path.exists(f'./user_data/{userName}/{chooseFile}/technicalTest/technicalTest.json'):
        with open(f'./user_data/{userName}/{chooseFile}/technicalTest/technicalTest.json' , encoding="utf-8") as f:
            inputJsons = json.load(f)
        wordFilePath = fillin(userName , chooseFile , inputJsons , "technicalTest")
        path = combineWord(wordFilePath , "technicalTest")
        combineWordPath.append(path)
        # docxConvert()
    
    combineWord(combineWordPath , "")
    docxConvert()
    
